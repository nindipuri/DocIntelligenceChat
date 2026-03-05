"""Parse uploaded documents into pages or logical sections."""

import logging
import re
import unicodedata
from pathlib import Path
from typing import Any

from pypdf import PdfReader

logger = logging.getLogger(__name__)


def _normalize_text(text: str) -> str:
    """Normalize unicode spaces and control chars so keyword matching works reliably."""
    if not text:
        return ""
    text = unicodedata.normalize("NFKC", text)
    text = re.sub(r"[\xa0\u200b\u200c\u200d\ufeff]+", " ", text)
    return " ".join(text.split())

# Chunk size for TXT splitting (characters)
TXT_CHUNK_MIN = 800
TXT_CHUNK_MAX = 1200


def _pages_to_dict_list(pages: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Ensure each item has 'page' and 'text' keys. Normalize text for retrieval."""
    return [
        {"page": p["page"], "text": _normalize_text(p.get("text", ""))}
        for p in pages
    ]


def load_pdf(file_path: Path) -> list[dict[str, Any]]:
    """
    Parse a PDF file into one entry per page.
    Uses PyMuPDF if available (better for tables/complex layouts), else pypdf.
    Returns list of {"page": 1-based index, "text": "..."}.
    """
    pages: list[dict[str, Any]] = []
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(file_path))
        for i in range(len(doc)):
            page = doc[i]
            text = page.get_text("text") or ""
            pages.append({"page": i + 1, "text": text.strip()})
        doc.close()
        logger.info("Parsed PDF %s with PyMuPDF: %d pages", file_path.name, len(pages))
    except ImportError:
        reader = PdfReader(str(file_path))
        for i, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            pages.append({"page": i, "text": text.strip()})
        logger.info("Parsed PDF %s with pypdf: %d pages", file_path.name, len(pages))
    return _pages_to_dict_list(pages)


def load_docx(file_path: Path) -> list[dict[str, Any]]:
    """
    Parse a DOCX file into sections (by paragraph groups / headings).
    Each section becomes a logical "page" (1-based).
    """
    try:
        from docx import Document
    except ImportError:
        raise ValueError("DOCX support requires python-docx; install with: pip install python-docx")

    doc = Document(str(file_path))
    sections: list[dict[str, Any]] = []
    current_chunks: list[str] = []
    page_num = 1

    def flush_section():
        nonlocal page_num
        if current_chunks:
            text = "\n\n".join(current_chunks).strip()
            if text:
                sections.append({"page": page_num, "text": text})
                page_num += 1
        current_chunks.clear()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Start a new section on headings (short lines or style)
        is_heading = len(text) < 100 or (para.style and "heading" in (para.style.name or "").lower())
        if is_heading and current_chunks:
            flush_section()
        current_chunks.append(text)

    flush_section()
    if not sections:
        # Fallback: one big section
        full = "\n\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        if full:
            sections = [{"page": 1, "text": full}]
    logger.info("Parsed DOCX %s: %d sections", file_path.name, len(sections))
    return _pages_to_dict_list(sections)


def load_txt(file_path: Path) -> list[dict[str, Any]]:
    """
    Parse a TXT file into chunks of roughly 800–1200 characters per "page".
    """
    text = file_path.read_text(encoding="utf-8", errors="replace")
    pages: list[dict[str, Any]] = []
    start = 0
    page_num = 1
    while start < len(text):
        # Prefer splitting at paragraph (double newline) or sentence end
        end = min(start + TXT_CHUNK_MAX, len(text))
        chunk = text[start:end]
        if end < len(text):
            # Try to break at last newline or sentence
            last_br = chunk.rfind("\n\n")
            last_dot = chunk.rfind(". ")
            break_at = max(last_br, last_dot)
            if break_at > TXT_CHUNK_MIN or break_at == -1:
                if break_at > 0:
                    chunk = text[start : start + break_at + 1]
                    end = start + break_at + 1
        chunk = chunk.strip()
        if chunk:
            pages.append({"page": page_num, "text": chunk})
            page_num += 1
        start = end
    if not pages and text.strip():
        pages = [{"page": 1, "text": text.strip()}]
    logger.info("Parsed TXT %s: %d chunks", file_path.name, len(pages))
    return _pages_to_dict_list(pages)


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def load_document(file_path: Path) -> list[dict[str, Any]]:
    """
    Load a document from disk and return a list of pages/sections.
    Each item: {"page": int (1-based), "text": str}.

    Supports: PDF, DOCX, TXT.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return load_pdf(path)
    if suffix == ".docx":
        return load_docx(path)
    if suffix == ".txt":
        return load_txt(path)
    raise ValueError(f"Unsupported format: {suffix}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}")
